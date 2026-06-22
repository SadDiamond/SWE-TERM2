from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document


ROOT = Path("/Users/gordonzhao/Documents/CPT/SWE-TERM2")
SOURCE = Path("/Users/gordonzhao/Downloads/Gordon_Zhao_Assessment_2_OOP_Report.docx")
OUTPUT = ROOT / "docs/report_revision/output/Gordon_Zhao_Assessment_2_OOP_Report_Refined.docx"
DIAGRAMS = ROOT / "tmp/report_diagrams_render"


def replace_oop_section(doc: Document) -> None:
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "1.2 Why OOP suits a game")
    end = next(i for i, p in enumerate(doc.paragraphs[start + 1 :], start + 1) if p.text.strip().startswith("1.3 "))
    body = [p for p in doc.paragraphs[start + 1 : end] if p.text.strip()]
    replacement = [
        "A game is a real-time system made from many objects that exist at the same time. The player stores movement and health state, each enemy stores its own combat state, weapons track cooldowns, and the arena director tracks floor progress. OOP suits this because each object can own its data and expose only the methods that other systems need.",
        "Encapsulation: PlayerController owns momentum, grounded state, dash charges and grapple state. Other scripts use public methods such as TakeDamage instead of directly changing movement variables. This matters because movement is the main mechanic and small unintended changes can create major physics bugs.",
        "Inheritance: Terminal, CybergrindShopStation and CybergrindWeaponReward inherit from Interactable. They share focus behaviour, prompt range and the OnInteract(PlayerController) contract, while each subclass implements a different result.",
        "Polymorphism: Gun resolves a hit, finds IDamageable and calls TakeDamage(float). The same call works for PlayerController, BasicEnemyAI and target objects, so weapon code does not need a separate branch for every target type.",
        "Abstraction and interface segregation: IGrappleMassTarget only exposes the mass category and pull behaviour required by the grapple. Damage and grapple response stay as separate contracts, so an object can implement either interface or both without inheriting from an unrelated class.",
        "Composition: Unity GameObjects combine components instead of relying on one deep inheritance tree. For example, a BasicEnemyAI object can use its own script together with a NavMeshAgent, colliders and renderers. This keeps navigation, physics and visuals replaceable without rewriting the whole enemy class.",
        "The current repository contains 68 C# scripts and about 35,407 lines under Assets/Scripts. The main classes are still too large and should eventually be split into smaller movement, weapon, AI-state and generation components. Even with that technical debt, the current object model is easier to debug than a single procedural loop with shared global state.",
    ]
    if len(body) != len(replacement):
        raise RuntimeError(f"Expected 7 paragraphs in section 1.2, found {len(body)}")
    for paragraph, text in zip(body, replacement):
        paragraph.text = text


def replace_media(docx_path: Path) -> None:
    mapping = {
        "word/media/image1.png": DIAGRAMS / "figure2_structure_chart.png",
        "word/media/image2.png": DIAGRAMS / "figure3_class_diagram.png",
        "word/media/image3.png": DIAGRAMS / "figure1_dfd.png",
    }
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        temp_path = Path(handle.name)
    try:
        with zipfile.ZipFile(docx_path, "r") as source_zip, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target_zip:
            for item in source_zip.infolist():
                data = mapping[item.filename].read_bytes() if item.filename in mapping else source_zip.read(item.filename)
                target_zip.writestr(item, data)
        shutil.move(temp_path, docx_path)
    finally:
        temp_path.unlink(missing_ok=True)


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    replace_oop_section(doc)
    doc.core_properties.title = "Software Engineering Assessment 2 - Project Report"
    doc.core_properties.subject = "Object-oriented Unity arena shooter project report"
    doc.save(OUTPUT)
    replace_media(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
